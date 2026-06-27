"""
DOCX 简历生成服务
将 AI 优化后的简历文本生成为格式化的 DOCX 文件
"""
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import re


def generate_resume_docx(resume_text: str, output_path: str):
    """将优化后的简历文本生成格式化的 DOCX 文件"""
    doc = Document()

    # 设置默认字体
    style = doc.styles["Normal"]
    font = style.font
    font.name = "微软雅黑"
    font.size = Pt(10.5)
    style.element.rPr.rFonts.set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia", "微软雅黑")

    # 设置页边距
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    lines = resume_text.strip().split("\n")

    # 先处理姓名作为标题
    name_found = False
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        if not line:
            i += 1
            continue

        # 判断是否是姓名行（第一行有效内容，短且不含特殊字符）
        if not name_found and len(line) < 20 and not any(
            kw in line for kw in ["【", "】", "联系", "电话", "@", "教育", "工作", "项目", "技能", "经验", "总结", "求职"]
        ):
            # 可能是姓名
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(line)
            run.font.size = Pt(18)
            run.font.bold = True
            run.font.name = "微软雅黑"
            run.element.rPr.rFonts.set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia", "微软雅黑")
            name_found = True
            i += 1
            continue

        # 判断是否是联系方式行
        if re.search(r"[\d+]{10,}|@", line) and len(line) < 60:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(line)
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(100, 100, 100)
            run.font.name = "微软雅黑"
            run.element.rPr.rFonts.set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia", "微软雅黑")
            i += 1
            continue

        # 判断是否是求职意向
        if "求职意向" in line or "目标岗位" in line or "应聘岗位" in line:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(line)
            run.font.size = Pt(10)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0, 82, 148)
            run.font.name = "微软雅黑"
            run.element.rPr.rFonts.set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia", "微软雅黑")
            i += 1
            continue

        # 判断是否是章节标题（以【】开头或包含特定关键词）
        is_section = False
        section_keywords = ["教育背景", "教育经历", "工作经历", "项目经验", "重点项目", "职业总结", "技术技能",
                            "专业技能", "个人优势", "自我评价", "证书", "语言能力", "求职意向"]
        for kw in section_keywords:
            if kw in line and len(line) < 30:
                is_section = True
                break

        if is_section or line.startswith("【") or line.startswith("##"):
            # 章节标题
            clean_title = line.replace("【", "").replace("】", "").replace("##", "").strip()
            # 添加分隔线
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run("━" * 40)
            run.font.size = Pt(6)
            run.font.color.rgb = RGBColor(200, 200, 200)

            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run(clean_title)
            run.font.size = Pt(12)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0, 82, 148)
            run.font.name = "微软雅黑"
            run.element.rPr.rFonts.set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia", "微软雅黑")
            i += 1
            continue

        # 判断是否是分隔线
        if set(line.strip()) <= {"-", "=", "━", "—", "─"}:
            i += 1
            continue

        # 普通内容行
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.5

        # 判断是否是列表项
        if line.startswith("-") or line.startswith("•") or line.startswith("·") or (len(line) > 1 and line[0].isdigit() and line[1] in ".、)）"):
            # 列表项，缩进
            p.paragraph_format.left_indent = Cm(0.5)
            text = line.lstrip("-•· 0123456789.、)） ")
            prefix = line[:len(line) - len(text)]
            run = p.add_run(prefix + text)
        else:
            run = p.add_run(line)

        run.font.size = Pt(10.5)
        run.font.name = "微软雅黑"
        run.element.rPr.rFonts.set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia", "微软雅黑")
        i += 1

    doc.save(output_path)
    return output_path